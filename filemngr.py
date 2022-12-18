import os
import cv2
import numpy as np
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image


class fmaneger():
    def __init__(self):
        self.MAINPATH = ''
        self.FILELIST= ''
        self.IMG = None
        self._IMGPROC = None
        self._IMGPROCINFO = {}
        self._MASK = None
        self._IMGNAME = ''
    def opendir(self):
        path = filedialog.askdirectory()
        filelist = sorted(os.listdir(path), key=len)
        for fichier in filelist:
            if not (fichier.endswith(".png") or fichier.endswith(".jpg")):
                filelist.remove(fichier)

        self.MAINPATH = path
        self.FILELIST = filelist
        self.IMG = None
        self.IMGPROC = None
        self.IMGPROCINFO = ()
    def openimg(self, path):
        img = cv2.imdecode(np.fromfile(self.MAINPATH+'\\'+path, dtype=np.uint8), cv2.IMREAD_COLOR)
        self._IMGNAME = path
        self.IMG = img
    def consider(self):
        try:
            if (self._IMGPROC == None):
                messagebox.showerror(
                    "Ошибка обработки изображения",
                    "Пожалуйста, выберете изображениe формата PNG или JPG")
        except:
            cv2.imshow('Consider menu', self._IMGPROC)
            cv2.waitKey(0)

    def getimgporc(self, img, info, mask):
        self._IMGPROC = img
        self._IMGPROCINFO = info
        self._MASK = mask

    def save(self):
        try:
            if (self._IMGPROC == None):
                messagebox.showerror(
                    "Ошибка сохранения",
                    "Пожалуйста, выберете изображениe формата PNG или JPG")
        except:
            path = filedialog.askdirectory()
            path += '/processing_result'
            try:
                os.mkdir(path)
            except:
                pass
            size = self._IMGPROCINFO[2].split('x')
            size = (int(size[0]), int(size[1]))
            _, im_buf_arr = cv2.imencode(".png", cv2.resize(self._IMGPROC,size))
            im_buf_arr.tofile(f'{path}/{self._IMGNAME}')
            Image.fromarray((self._MASK* 255).astype(np.uint8)).save(f'{path}/mask_{self._IMGNAME}')
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = "Calibri"
            style.font.size = Pt(16)

            doc.add_picture(path + f'/{self._IMGNAME}', width=Inches(5))
            doc.add_paragraph(f"number of bad blood cells: {self._IMGPROCINFO[0]}")
            doc.add_paragraph(f"number of good blood cells: {self._IMGPROCINFO[1]}")
            doc.add_paragraph(f"Image size: {self._IMGPROCINFO[2]}")
            doc.add_paragraph(f"Date and time image processing: {self._IMGPROCINFO[3]}")
            doc.save (f"{path}/report_{self._IMGNAME[:-4]}.doc")

        
